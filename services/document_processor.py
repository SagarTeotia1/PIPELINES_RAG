import os
import uuid
import logging
from typing import List, Dict, Any, Optional
from PyPDF2 import PdfReader
from docx import Document
from config import Config

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.txt']
    
    def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process a document and return metadata and chunks"""
        try:
            # Validate file extension
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Extract text based on file type
            if file_ext == '.pdf':
                text = self._extract_pdf_text(file_path)
            elif file_ext == '.docx':
                text = self._extract_docx_text(file_path)
            elif file_ext == '.txt':
                text = self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Validate text
            if not text or len(text.strip()) < 10:
                raise ValueError("Document contains insufficient text content")
            
            # Create document metadata
            document_id = str(uuid.uuid4())
            document_metadata = {
                'document_id': document_id,
                'filename': filename,
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'file_extension': file_ext,
                'total_characters': len(text),
                'total_pages': self._count_pages(file_path, file_ext),
                'processing_status': 'completed',
                'chunks_count': 0
            }
            
            # Chunk the text
            chunks = self._chunk_text(text, document_id, filename)
            document_metadata['chunks_count'] = len(chunks)
            
            return {
                'metadata': document_metadata,
                'chunks': chunks
            }
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            reader = PdfReader(file_path)
            text = ""
            
            for page_num, page in enumerate(reader.pages):
                if page_num >= Config.MAX_PAGES_PER_DOCUMENT:
                    logger.warning(f"PDF exceeds maximum pages limit. Processing first {Config.MAX_PAGES_PER_DOCUMENT} pages.")
                    break
                
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            raise
    
    def _count_pages(self, file_path: str, file_ext: str) -> int:
        """Count pages in a document"""
        try:
            if file_ext == '.pdf':
                reader = PdfReader(file_path)
                return len(reader.pages)
            elif file_ext == '.docx':
                doc = Document(file_path)
                # DOCX doesn't have a direct page count, estimate based on content
                return max(1, len(doc.paragraphs) // 50)
            elif file_ext == '.txt':
                # Estimate pages for text files (assuming 2000 characters per page)
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                return max(1, len(content) // 2000)
            else:
                return 1
                
        except Exception as e:
            logger.error(f"Error counting pages: {e}")
            return 1
    
    def _chunk_text(self, text: str, document_id: str, filename: str) -> List[Dict[str, Any]]:
        """Split text into chunks"""
        try:
            chunks = []
            words = text.split()
            chunk_size = Config.CHUNK_SIZE
            overlap = Config.CHUNK_OVERLAP
            
            for i in range(0, len(words), chunk_size - overlap):
                chunk_words = words[i:i + chunk_size]
                chunk_text = " ".join(chunk_words)
                
                if len(chunk_text.strip()) < 50:  # Skip very short chunks
                    continue
                
                chunk_id = f"{document_id}_chunk_{i//chunk_size}"
                
                chunk_data = {
                    'chunk_id': chunk_id,
                    'document_id': document_id,
                    'filename': filename,
                    'text': chunk_text,
                    'chunk_index': i // chunk_size,
                    'word_count': len(chunk_words),
                    'start_word_index': i,
                    'end_word_index': min(i + chunk_size, len(words))
                }
                
                chunks.append(chunk_data)
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error chunking text: {e}")
            raise
    
    def validate_file(self, filename: str, file_size: int) -> bool:
        """Validate uploaded file"""
        try:
            # Check file extension
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext not in self.supported_extensions:
                return False
            
            # Check file size (max 50MB)
            max_size = 50 * 1024 * 1024  # 50MB
            if file_size > max_size:
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error validating file: {e}")
            return False
